"""
cv_scanner.py - Walk a folder and extract raw text from PDF and DOCX files.
Returns (file_path, text, md5_hash) tuples via a generator.
"""

import os
import hashlib
import logging

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def compute_md5(file_path: str) -> str:
    h = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_text_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
    except Exception as e:
        logger.warning(f"PDF extraction failed for {file_path}: {e}")
        return ""


def extract_text_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"DOCX extraction failed for {file_path}: {e}")
        return ""


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_docx(file_path)
    return ""


def scan_folder(folder_path: str, known_hashes: set = None):
    """
    Generator that yields dicts:
    {
        "file_path": str,
        "file_name": str,
        "text": str,
        "md5": str,
        "error": str or None,
        "already_known": bool   # True = hash in DB, text extraction skipped
    }
    If known_hashes is provided, files whose MD5 is already in the set
    are yielded immediately (no text extraction) with already_known=True.
    """
    if not os.path.isdir(folder_path):
        raise ValueError(f"Folder not found: {folder_path}")

    for root, _dirs, files in os.walk(folder_path):
        for fname in files:
            ext = os.path.splitext(fname)[1].lower()
            if ext not in SUPPORTED_EXTENSIONS:
                continue

            file_path = os.path.join(root, fname)
            try:
                md5 = compute_md5(file_path)

                # Skip text extraction for already-known files
                if known_hashes is not None and md5 in known_hashes:
                    yield {
                        "file_path": file_path,
                        "file_name": fname,
                        "text": "",
                        "md5": md5,
                        "error": None,
                        "already_known": True,
                    }
                    continue

                text = extract_text(file_path)
                yield {
                    "file_path": file_path,
                    "file_name": fname,
                    "text": text,
                    "md5": md5,
                    "error": None,
                    "already_known": False,
                }
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                yield {
                    "file_path": file_path,
                    "file_name": fname,
                    "text": "",
                    "md5": "",
                    "error": str(e),
                    "already_known": False,
                }


def count_files(folder_path: str) -> int:
    """Count supported CV files in folder (for progress reporting)."""
    count = 0
    for root, _dirs, files in os.walk(folder_path):
        for fname in files:
            if os.path.splitext(fname)[1].lower() in SUPPORTED_EXTENSIONS:
                count += 1
    return count
